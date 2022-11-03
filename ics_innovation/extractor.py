# import textract
import csv
import sys
from glob import glob
import docx
import fitz
import ntpath
import os, sys
import docx2txt
import traceback

flat_list = list()

class exceptionMessage(Exception):
    def __init__(self, func_name, message):
        self.message = message
        if(type(message) == str):
            self.mesage = "no keyword found"
        else:
            _, _, exception_traceback = sys.exc_info()
            file_name = os.path.split(exception_traceback.tb_frame.f_code.co_filename)[1]
            self.message = f"ERROR | {func_name} | {message} | {file_name} | {exception_traceback.tb_lineno}"
            self.message = f"ERROR: {message} at {func_name}@{exception_traceback.tb_lineno} in {file_name}\n"
            
    def __str__(self):
        return "{0}".format(self.message)

def get_config_folder_path():
    return os.path.join(os.getcwd(), "config")


def write_to_text(file_name, text):
    f = open("extracted_file_text.txt", "a", encoding="utf-8")
    f.write("*"*100)
    f.write("\n")
    f.write(file_name)
    f.write("\n")
    f.write("-"*100)
    f.write("\n")
    f.write(text)
    f.write("\n")
    f.write("#"*100)
    f.write("\n")
    f.close() 

def get_text_from_file(file):
    '''
        extract text from file
    '''
    actual_fname = os.path.splitext(file)
    isPDF = False
    extracted_dates_dict = {}
    rule_id = "1"
    block=""
    if actual_fname[1] == '.doc':
        # logging.info("Loaded the doc file")
        doc = open(file, 'r', encoding='utf-8', errors='ignore')
        block = doc.read()
        # block2 = []
    elif actual_fname[1] == '.docx':
        # logging.info("Loaded the docx file")
        fulltxt = docx2txt.process(file)
        # doc = docx.Document(file)
        # fulltxt = ''
        # if doc.tables:
        #     tables = doc.tables
        #     for table in tables:
        #         for row in table.rows:
        #             # print(row)
        #             for cell in row.cells:
        #                 # print(cell.text)
        #                 for para in cell.paragraphs:
        #                     # print("paratext",para.text)
        #                     fulltxt +=para.text
        # #header footer extraction starts here
        # if doc.sections[0]:
        #     # >>>> header and footer reference
        #     sec = doc.sections[0]
        #     header = sec.header
        #     footer = sec.footer
        #     for para1 in header.paragraphs:
        #         # print("para1",para1.text)
        #         fulltxt += para1.text
        #     for para2 in footer.paragraphs:
        #         # print("para2",para2.text)
        #         fulltxt += para2.text
        # #---------------------- header footer extraction ends here ---------
        # #******************* body extraction starts here **************
        # if doc.paragraphs:

        #     for para in doc.paragraphs:
        #         fulltxt += para.text

        block = fulltxt
        # block2 = []
        #------------------ body extraction ends here -----------------------
        # print(.block)
        # (doc.paragraphs)
    elif actual_fname[1] == '.pdf':
        # logging.info("Loaded the pdf file")
        isPDF = True
        doc = fitz.open(file)
        pipdf2 = fitz.open(file)
        page = pipdf2.load_page(0)
        page2 = pipdf2.load_page(len(pipdf2)-1)
        block = page.get_text("blocks")
        pdf_content = page.get_text()
        if(len(pipdf2)>1):
            block2 = page2.get_text("blocks")
            page2 = pipdf2.load_page(len(pipdf2)-1)
            block = block + page2.get_text("blocks")
            pdf_content = pdf_content + page2.get_text()

    
    elif actual_fname[1] == '.msg':
        with extract_msg.Message(file) as msg:
            block = msg.body
            msg_subject = msg.subject
        # print(msg_subject)
    file_content =  pdf_content if isPDF else block
    write_to_text(file, file_content)
    return file_content

def get_title_from_file(file):
    '''
        extract title from file
    '''
    actual_fname = os.path.splitext(file)
    # isPDF = False
    # extracted_dates_dict = {}
    # rule_id = "1"
    title=""
    if actual_fname[1] == '.doc':
        # logging.info("Loaded the doc file")
        doc = open(file, 'r', encoding='cp1252', errors='ignore')
        block = doc.read()
        title=block.split('\n')[3]
    elif actual_fname[1] == '.docx':
        # logging.info("Loaded the docx file")
        doc = docx.Document(file)
        title=doc.core_properties.title          #------------------ body extraction ends here -----------------------
        # print(.block)
        # (doc.paragraphs)
    elif actual_fname[1] == '.pdf':
        # logging.info("Loaded the pdf file")
        pipdf2 = fitz.open(file)
        title=pipdf2.metadata['title']
    # file_content =  pdf_content if isPDF else block
    else:
        data=textract.process(file).decode('utf-8')
        title=data.split("\n")[0]
    # file_content =  pdf_content if isPDF else block
    return title


    
def get_fulltext_from_pdf(file):
    actual_fname = os.path.splitext(file)
    if actual_fname[1] == '.pdf':
        pages=""
        doc = fitz.open(file)
        for page in range(0,len(doc)):
            pages=pages+doc.load_page(page).get_text()
        total_data = pages 
        # print(total_data)
        write_to_text(file, total_data)
        return total_data 
    else:
        data=get_text_from_file(file)
        return data

def get_pagetext_from_pdf(file,page_no):
    actual_fname = os.path.splitext(file)
    if actual_fname[1] == '.pdf':
        pages=""
        doc = fitz.open(file)
        if len(doc)>1 and (len(doc)-1 != page_no):
            page = doc.load_page(page_no)
            block = page.get_text("blocks")
            pdf_content = page.get_text()
            # print(pdf_content)
            write_to_text(file, pdf_content)
            return pdf_content
        else:
            # print("No text")
            write_to_text(file, "No text")
            return ""
        #page = doc.load_page(page_no)
        #block = page.get_text("blocks")
        #pdf_content = page.get_text()
        #return pdf_content

def get_first_page_text(file):
    actual_fname = os.path.splitext(file)
    if actual_fname[1] == '.pdf':
        pages=""
        doc = fitz.open(file)
        page = doc.load_page(0)
        total_data = page.get_text()
        # print(total_data)
        write_to_text(file, total_data)
        return total_data 
    else:
        data=get_text_from_file(file)
        return data

def get_first_page_text_percentage(file, percentage):
    actual_fname = os.path.splitext(file)
    if actual_fname[1] == '.pdf':
        pages=""
        doc = fitz.open(file)
        page = doc.load_page(0)
        total_data = page.get_text()
        # print(total_data)
        write_to_text(file, total_data)
        return total_data
    else:
        data=get_text_from_file_percentage(file, percentage)
        return data


def get_text_from_file_percentage(file, percentage):
    #f = open(input_file, 'r', encoding='utf-8', errors='ignore')
    #text = f.read()
    # actual_fname = os.path.splitext(file)
    # if actual_fname[1] == '.docx':
    text = file
    #print("@@@@@@@@@@@@@@@@@@",text)
    word_list = text.split()
    len_words = len(word_list)
    per = percentage / 100
    per_res = round(per * len_words)
    words = word_list[0:per_res]
    result = ' '.join([str(word) for word in words])
    # print("!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!", result)
    #print("!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!!")
    return result

def get_only_text_docx(file):
    actual_fname = os.path.splitext(file)
    data=""
    if actual_fname[1] == '.docx':
        doc = docx.Document(file)
        fulltxt = ''
        if doc.sections[0]:
            # >>>> header and footer reference
            sec = doc.sections[0]
            header = sec.header
            footer = sec.footer
            # for para in doc.paragraphs:
            #     for run in para.runs:
            #         if run.bold :
            #             fulltxt +=run.text+"\n"
            # for para in doc.paragraphs:
            #     fulltxt += para.text+"\n"

            for para1 in header.paragraphs:
                # print("para1",para1.text)
                fulltxt += para1.text+"\n"
            for para2 in footer.paragraphs:
                # print("para2",para2.text)
                fulltxt += para2.text+"\n"
        if doc.paragraphs:

            for para in doc.paragraphs:
                fulltxt += para.text+"\n"
        # data= docx2txt.process(file)
        data=fulltxt
        # print(data)
        write_to_text(file, data)
        return data
    elif actual_fname[1] == '.pdf':
        return get_text_from_file_with_any_language(file)
    else:
        data=get_text_from_file(file)
        return data

def get_text_from_file_with_any_language(file):
    actual_fname = os.path.splitext(file)
    data=""
    if actual_fname[1] == '.pdf':
        # print("pdf file")
        doc = fitz.open(file)
    
        pages=""
        for page in range(0,len(doc)):
            pages=pages+doc.getPageText(page)
        # print(pages)
        write_to_text(file, pages)
        return  pages

    else:
        data=get_text_from_file(file)
        return data

def get_tables_from_doccumet(file):
    try:
        actual_fname = os.path.splitext(file)
        json_={}
        json_table_list=[]
        if actual_fname[1] == '.docx' or actual_fname[1]  == '.docm':
            doc = docx.Document(file)
            if doc.tables:
                tables = doc.tables
                for table in range(len(tables)):
                    complete_rows=[]
                    col_valss=[]
    #                 indexes=[]
                    rows=tables[table].rows
                    for row in range(len(rows)):
                        row_vals=[]
                        for cell in rows[row].cells:
                            for para in cell.paragraphs:     
                                if row==0:
                                    col_valss.append(para.text)
                                else:
    #                                 indexes.append(row-1)
                                    row_vals.append(para.text)
                        if row!=0:
                            extend_length =len(col_valss)-len(row_vals)
                            if extend_length>0:
                                row_vals.extend(extend_length * ['nan'])
                            else:
                                col_valss.extend(-extend_length * [' '])
                            complete_rows.append(row_vals)
                    table_=pd.DataFrame(complete_rows,columns=col_valss)
                    json_table_list.append(table_) 
            return json_table_list
        if actual_fname[1]  == '.pdf':
            # print("PDF$$$$$$$$$$$$$$$$$$")
            pipdf2 = fitz.open(file)
            tables=[]
            # print("len(pipdf2)",len(pipdf2))
            no_of_pages=len(pipdf2)
            if no_of_pages < 4:
                tables = tabula.read_pdf(file, pages="1-" + str(no_of_pages))
            else:
                tables = tabula.read_pdf(file, pages="1-4")
            return tables
    except Exception as x:
        print("exception:get_tables_from_doccumet",x)
        return []
        raise exceptionMessage("get_tables_from_doccumet", x)

def get_tables_from_document(file):
    try:
        actual_fname = os.path.splitext(file)
        json_={}
        json_table_list=[]
        if actual_fname[1] == '.docx' or actual_fname[1]  == '.docm':
            doc = docx.Document(file)
            if doc.tables:
                # print(doc.tables)
                tables = doc.tables
                for table in range(len(tables)):
                    # print("index-->",table)
                    complete_rows=[]
                    col_valss=[]
    #                 indexes=[]
                    rows=tables[table].rows
                    for row in range(len(rows)):
                        row_vals=[]
                        for cell in rows[row].cells:
                            for para in cell.paragraphs:     
                                if row==0:
                                    col_valss.append(para.text)
                                else:
    #                                 indexes.append(row-1)
                                    row_vals.append(para.text)
                        if row!=0:
                            extend_length =len(col_valss)-len(row_vals)
                            if extend_length>0:
                                row_vals.extend(extend_length * ['nan'])
                            else:
                                col_valss.extend(-extend_length * [' '])
                            complete_rows.append(row_vals)
                    table_=pd.DataFrame(complete_rows,columns=col_valss)
                    json_table_list.append(table_) 
            return json_table_list
        if actual_fname[1]  == '.pdf':
            pipdf2 = fitz.open(file)
            tables=[]
            # print("len(pipdf2)",len(pipdf2))
            no_of_pages=len(pipdf2)
            if no_of_pages<4:
                tables=tabula.read_pdf(file,pages="1-"+str(no_of_pages), pandas_options={'header':None})
            else:
                tables=tabula.read_pdf(file,pages="1-4", pandas_options={'header':None})
            return tables
    except Exception as x:
        raise exceptionMessage("get_tables_from_document", x)

def nest_list_to_flat(input_list: list) -> list:
    """
        Convert Nested list to flat list
    """
    try:
        for in_list in input_list:
            if type(in_list) == list:
                nest_list_to_flat(in_list)
            else:
                flat_list.append(in_list)
        return flat_list
    except Exception as x:
        raise exceptionMessage("nest_list_to_flat", x)